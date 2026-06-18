import json
import re
import datetime
import io
import os

from openai import OpenAI
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.ai_content_detector import (
    _PRIMARY,
    _HEADER_FILL,
    _HEADER_TEXT,
    _ZEBRA_FILL,
    _MUTED,
    _FOOTER,
    _shade_cell,
    _set_cell,
    _set_table_borders,
    _set_col_widths,
    _add_header_row,
    _verdict_palette,
    _add_styled_heading,
    _add_section_if_text,
    _fmt_pct,
)

from src.databricks_client import get_openai_client, MODEL_GPT4O

OPENAI_MODEL = MODEL_GPT4O

_client = None


def _get_client():
    global _client
    if _client is None:
        _client = get_openai_client()
    return _client


def _to_float(v, default=0.0):
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def _to_bool(v, default=False):
    if isinstance(v, bool):
        return v
    if isinstance(v, str):
        return v.strip().lower() in ("true", "yes", "1")
    return default


def _extract_json_object(raw: str) -> str:
    s = raw.strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s)
        s = re.sub(r"\s*```\s*$", "", s)
    start = s.find("{")
    if start < 0:
        return s
    depth = 0
    in_str = False
    esc = False
    for i in range(start, len(s)):
        c = s[i]
        if esc:
            esc = False
            continue
        if c == "\\":
            esc = True
            continue
        if c == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return s[start:i + 1]
    return s[start:]


def _repair_inner_quotes(s: str) -> str:
    out = []
    in_string = False
    esc = False
    i = 0
    n = len(s)
    while i < n:
        c = s[i]
        if esc:
            out.append(c)
            esc = False
            i += 1
            continue
        if c == "\\":
            out.append(c)
            esc = True
            i += 1
            continue
        if c == '"':
            if not in_string:
                in_string = True
                out.append(c)
            else:
                j = i + 1
                while j < n and s[j] in " \t\r\n":
                    j += 1
                if j >= n or s[j] in ",:}]":
                    in_string = False
                    out.append(c)
                else:
                    out.append('\\"')
            i += 1
            continue
        out.append(c)
        i += 1
    return "".join(out)


def _safe_json_parse(raw: str) -> dict:
    payload = _extract_json_object(raw)
    try:
        return json.loads(payload)
    except json.JSONDecodeError:
        pass
    try:
        return json.loads(payload, strict=False)
    except json.JSONDecodeError:
        pass
    no_trailing = re.sub(r",(\s*[}\]])", r"\1", payload)
    try:
        return json.loads(no_trailing, strict=False)
    except json.JSONDecodeError:
        pass
    repaired = _repair_inner_quotes(no_trailing)
    return json.loads(repaired, strict=False)


SYSTEM_PROMPT = """You are an LLM-as-a-Judge reviewing an AI Content Detector result.

You will receive:
1. The original document text.
2. A readable summary of the detector's assessment.

Your job has TWO parts:

PART A — Independent verdict (second-opinion):
Form your OWN judgment about whether the document is human-written, AI-assisted, or AI-generated, without anchoring on the detector's verdict. Then compare to the detector and explicitly note agreement or disagreement.

PART B — Reasoning audit:
Audit only user-facing, material detector reasoning:
- The overall verdict and confidence.
- Listed AI-writing markers.
- Listed human-writing markers.
- Flagged source-document sentences.
- Section assessments only when they contradict quoted source text.

Do NOT audit internal report fields or generated report prose as source evidence:
- Do not audit summary_pct, sections arrays, human_markers arrays, ai_markers arrays, executive_summary, methodology_notes, or any JSON field name.
- Do not treat the detector's generated summary as the uploaded document.
- Do not flag estimated percentages merely because they are estimates. Only flag them if they contradict the verdict or sentence-level evidence.
- Do not output braces, JSON snippets, raw field names, or placeholder text like "[ ... ]".

Flag:
- Markers whose quoted evidence does not appear in the original document.
- Section or sentence verdicts that contradict the cited original-document evidence.
- Internal contradictions (e.g., 100% human but lists AI markers).
- Confidence levels that don't match the underlying evidence quality.

For each audit item, set severity: "high" (likely fabricated or contradictory), "medium" (weakly supported), "low" (minor wording issue).
Only include real issues in reasoning_audit. If there are no material issues, return an empty array. Do not include "verified" rows.

Final overall_judgment:
- "Trustworthy" — verdict and reasoning both well-grounded.
- "Suspect" — verdict plausible but reasoning has unsupported claims.
- "Unreliable" — verdict likely wrong OR reasoning has high-severity fabrications.

OUTPUT: valid JSON only, exactly this schema. No prose before or after.
{
  "independent_verdict": "Human Written" | "AI-assisted" | "AI-generated",
  "independent_confidence": 0.0,
  "independent_reasoning": "2-3 sentence justification grounded in specific text",
  "agrees_with_detector": true,
  "disagreement_explanation": "empty string if agrees, otherwise explain the difference",
  "reasoning_audit": [
    {
      "detector_claim": "plain-English detector claim, no JSON or braces",
      "verifiable": true,
      "issue": "what is wrong and why it matters",
      "severity": "high" | "medium" | "low"
    }
  ],
  "overall_judgment": "Trustworthy" | "Suspect" | "Unreliable",
  "recommended_action": "what the human reviewer should do"
}"""


def _fmt_conf(value) -> str:
    return f"{round(_to_float(value) * 100, 1)}%"


def _format_detector_for_judge(detector_result: dict) -> str:
    lines = [
        f"Overall verdict: {detector_result.get('doc_verdict', 'Unknown')}",
        f"Confidence: {_fmt_conf(detector_result.get('doc_confidence', 0.0))}",
    ]

    summary = detector_result.get("executive_summary")
    if summary:
        lines.extend(["", f"Detector summary: {summary}"])

    sections = detector_result.get("sections", []) or []
    if sections:
        lines.append("")
        lines.append("Section assessments:")
        for section in sections[:8]:
            lines.append(
                "- "
                f"{section.get('name', 'Section')}: "
                f"{section.get('verdict', 'Unknown')} "
                f"({_fmt_conf(section.get('confidence', 0.0))})"
                f" — {section.get('summary', '')}"
            )

    flagged = [
        sentence for sentence in (detector_result.get("sentences", []) or [])
        if str(sentence.get("verdict", "")).lower() != "human"
    ]
    if flagged:
        lines.append("")
        lines.append("Flagged source-document sentences:")
        for sentence in flagged[:12]:
            lines.append(
                "- "
                f"{sentence.get('id', 'sentence')}: "
                f"{sentence.get('verdict', 'Unknown')} "
                f"({_fmt_conf(sentence.get('score', 0.0))}) — "
                f"\"{sentence.get('text', '')}\" "
                f"Reason: {sentence.get('reason', '')}"
            )

    human_markers = detector_result.get("human_markers", []) or []
    if human_markers:
        lines.append("")
        lines.append("Human-writing markers claimed:")
        for marker in human_markers[:8]:
            lines.append(f"- {marker.get('type', 'Marker')}: \"{marker.get('evidence', '')}\"")

    ai_markers = detector_result.get("ai_markers", []) or []
    if ai_markers:
        lines.append("")
        lines.append("AI-writing markers claimed:")
        for marker in ai_markers[:8]:
            lines.append(
                f"- {marker.get('type', 'Marker')}: "
                f"\"{marker.get('evidence', '')}\" "
                f"({marker.get('location', 'location not specified')})"
            )

    return "\n".join(lines)


def judge_detector_output(narrative_text: str, detector_result: dict) -> dict:
    user_msg = (
        f"ORIGINAL DOCUMENT:\n{narrative_text}\n\n"
        f"DETECTOR ASSESSMENT:\n{_format_detector_for_judge(detector_result)}"
    )

    raw = ""
    try:
        resp = _get_client().chat.completions.create(
            model=OPENAI_MODEL,
            max_tokens=16000,
            temperature=0.0,
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_msg},
            ],
        )
        raw = (resp.choices[0].message.content or "").strip()
        return _safe_json_parse(raw)
    except json.JSONDecodeError as e:
        return {"error": f"Judge returned malformed JSON: {e}", "raw": raw[:1000]}
    except Exception as e:
        return {"error": f"Judge AI call failed: {e}"}


_MATERIAL_SEVERITIES = {"high", "medium", "low"}
_INTERNAL_AUDIT_FIELDS = (
    "summary_pct",
    "sections",
    "human_markers",
    "ai_markers",
    "executive_summary",
    "methodology_notes",
)


def _clean_audit_text(value) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    return text.translate(str.maketrans({
        "{": "",
        "}": "",
        "[": "",
        "]": "",
    }))


def _is_internal_audit_claim(claim: str) -> bool:
    normalized = claim.strip().lower()
    return any(
        normalized == field or normalized.startswith(f"{field}:")
        for field in _INTERNAL_AUDIT_FIELDS
    )


def _material_audit_items(audit: list) -> list:
    items = []
    for item in audit or []:
        severity = str(item.get("severity", "")).strip().lower()
        claim = _clean_audit_text(item.get("detector_claim", ""))
        issue = _clean_audit_text(item.get("issue", ""))
        if severity not in _MATERIAL_SEVERITIES:
            continue
        if _is_internal_audit_claim(claim):
            continue
        if issue.lower() in {"verified", "none", "no issue", "no issues"}:
            continue
        cleaned = dict(item)
        cleaned["severity"] = severity
        cleaned["detector_claim"] = claim
        cleaned["issue"] = issue
        items.append(cleaned)

    order = {"high": 0, "medium": 1, "low": 2}
    return sorted(items, key=lambda x: order.get(str(x.get("severity", "")).lower(), 3))


def generate_judge_report(judge_result: dict, detector_result: dict, source_name: str = "Document") -> str:
    if "error" in judge_result:
        return f"### LLM Judge Audit Report\n\n**Error:** {judge_result['error']}\n"

    today = datetime.datetime.now().strftime("%B %d, %Y")
    indep_verdict = judge_result.get("independent_verdict", "Unknown")
    indep_conf = _to_float(judge_result.get("independent_confidence", 0.0))
    detector_verdict = detector_result.get("doc_verdict", "Unknown")
    detector_conf = _to_float(detector_result.get("doc_confidence", 0.0))
    agrees = _to_bool(judge_result.get("agrees_with_detector", False))
    overall = judge_result.get("overall_judgment", "Unknown")

    md = "### LLM Judge Audit Report\n"
    md += "**Audit of AI Content Detector Output**\n\n"
    md += f"**Document:** {source_name}\n"
    md += f"**Date Audited:** {today}\n"
    md += f"**Verdict Reliability:** {overall}\n\n"

    md += "#### Verdict Comparison\n\n"
    md += "| Source | Verdict | Confidence |\n"
    md += "|---|---|---|\n"
    md += f"| Detector | {detector_verdict} | {round(detector_conf*100,1)}% |\n"
    md += f"| Judge (independent) | {indep_verdict} | {round(indep_conf*100,1)}% |\n"
    md += f"| Agreement | {'Yes' if agrees else 'No'} | — |\n\n"

    if not agrees:
        md += "**Disagreement Explanation:** "
        md += judge_result.get("disagreement_explanation", "") + "\n\n"

    md += "#### Independent Reasoning\n"
    md += judge_result.get("independent_reasoning", "") + "\n\n"

    audit = _material_audit_items(judge_result.get("reasoning_audit", []) or [])
    md += "#### Judge Findings\n\n"
    if audit:
        md += "| Severity | Finding | Why it matters |\n"
        md += "|---|---|---|\n"
        for item in audit:
            claim = (item.get("detector_claim", "") or "").replace("|", "\\|")
            issue = (item.get("issue", "") or "").replace("|", "\\|")
            sev = (item.get("severity", "") or "").upper()
            md += f"| {sev} | {claim} | {issue} |\n"
        md += "\n"
    else:
        md += "No material issues found in the detector reasoning.\n\n"

    high = [a for a in audit if str(a.get("severity", "")).lower() == "high"]
    if high:
        md += f"**High-severity issues found:** {len(high)}. Recommend human re-review of the detector output before relying on its verdict.\n\n"

    md += "#### Recommended Action\n"
    md += judge_result.get("recommended_action", "") + "\n\n"
    md += "*This audit complements (does not replace) human review.*\n"

    return md.replace("$", "\\$")


# ============================================================
# Word document generator for the judge audit. Same palette /
# helpers as the detector report so both .docx files visually
# match. The LLM only fills in `judge_result`; the design is
# fixed below.
# ============================================================

_JUDGMENT_COLORS = {
    "trustworthy": ("E2EFDA", "375623"),
    "suspect": ("FFF2CC", "806000"),
    "unreliable": ("FCE4D6", "843C0C"),
}

_SEVERITY_COLORS = {
    "high": ("FCE4D6", "843C0C"),
    "medium": ("FFF2CC", "806000"),
    "low": ("F2F2F2", "595959"),
    "none": ("E2EFDA", "375623"),
}


def _judgment_palette(judgment: str):
    return _JUDGMENT_COLORS.get((judgment or "").strip().lower(), ("F2F2F2", _MUTED))


def _severity_palette(severity: str):
    return _SEVERITY_COLORS.get((severity or "").strip().lower(), ("F2F2F2", _MUTED))


def generate_judge_docx(judge_result: dict, detector_result: dict, source_name: str = "Document") -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    today = datetime.datetime.now().strftime("%B %d, %Y")

    if "error" in judge_result:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("LLM Judge Audit Report")
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor.from_string(_PRIMARY)
        doc.add_paragraph(f"Error: {judge_result['error']}")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    indep_verdict = judge_result.get("independent_verdict", "Unknown")
    indep_conf = _to_float(judge_result.get("independent_confidence", 0.0))
    detector_verdict = detector_result.get("doc_verdict", "Unknown")
    detector_conf = _to_float(detector_result.get("doc_confidence", 0.0))
    agrees = _to_bool(judge_result.get("agrees_with_detector", False))
    overall = judge_result.get("overall_judgment", "Unknown")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LLM Judge Audit Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(_PRIMARY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Audit of AI Content Detector Output")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(_MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Document: {source_name}    |    Date Audited: {today}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(_MUTED)

    _add_styled_heading(doc, "Verdict Reliability", 1)
    j_fill, j_text = _judgment_palette(overall)
    badge = doc.add_paragraph()
    r1 = badge.add_run("Reliability: ")
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = badge.add_run(overall)
    r2.bold = True
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor.from_string(j_text)

    _add_styled_heading(doc, "Verdict Comparison", 2)
    tbl = doc.add_table(rows=4, cols=3)
    _add_header_row(tbl, ["Source", "Verdict", "Confidence"],
                    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER])

    det_fill, det_text = _verdict_palette(detector_verdict)
    indep_fill, indep_text = _verdict_palette(indep_verdict)

    _set_cell(tbl.rows[1].cells[0], "Detector", size=10, bold=True)
    _set_cell(tbl.rows[1].cells[1], detector_verdict, size=10, bold=True, color_hex=det_text)
    _shade_cell(tbl.rows[1].cells[1], det_fill)
    _set_cell(tbl.rows[1].cells[2], _fmt_pct(detector_conf), size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    _set_cell(tbl.rows[2].cells[0], "Judge (independent)", size=10, bold=True)
    _set_cell(tbl.rows[2].cells[1], indep_verdict, size=10, bold=True, color_hex=indep_text)
    _shade_cell(tbl.rows[2].cells[1], indep_fill)
    _set_cell(tbl.rows[2].cells[2], _fmt_pct(indep_conf), size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    agree_str = "Yes" if agrees else "No"
    agree_fill, agree_text = (_SEVERITY_COLORS["none"] if agrees else _SEVERITY_COLORS["high"])
    _set_cell(tbl.rows[3].cells[0], "Agreement", size=10, bold=True)
    _set_cell(tbl.rows[3].cells[1], agree_str, size=10, bold=True, color_hex=agree_text,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(tbl.rows[3].cells[1], agree_fill)
    _set_cell(tbl.rows[3].cells[2], "—", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [4.5, 6.5, 4.0])

    if not agrees:
        _add_section_if_text(doc, "Disagreement Explanation", 2,
                             judge_result.get("disagreement_explanation", ""))

    _add_section_if_text(doc, "Independent Reasoning", 2,
                         judge_result.get("independent_reasoning", ""))

    audit = _material_audit_items(judge_result.get("reasoning_audit", []) or [])
    _add_styled_heading(doc, "Judge Findings", 2)
    if audit:
        tbl = doc.add_table(rows=1 + len(audit), cols=3)
        _add_header_row(tbl, ["Severity", "Finding", "Why It Matters"],
                        aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                                WD_ALIGN_PARAGRAPH.LEFT])
        for ri, item in enumerate(audit, start=1):
            sev = (item.get("severity", "") or "").strip().lower()
            sev_fill, sev_text = _severity_palette(sev)
            row = tbl.rows[ri]
            _set_cell(row.cells[0], sev.upper() or "—", size=10, bold=True,
                      color_hex=sev_text, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade_cell(row.cells[0], sev_fill)
            _set_cell(row.cells[1], item.get("detector_claim", "") or "", size=10)
            _set_cell(row.cells[2], item.get("issue", "") or "", size=10)
            if ri % 2 == 0:
                for ci in (1, 2):
                    _shade_cell(row.cells[ci], _ZEBRA_FILL)
        _set_table_borders(tbl)
        _set_col_widths(tbl, [2.0, 6.5, 7.5])

        high_count = sum(1 for a in audit if str(a.get("severity", "")).lower() == "high")
        if high_count:
            p = doc.add_paragraph()
            r = p.add_run(
                f"High-severity issues found: {high_count}. "
                "Recommend human re-review of the detector output before relying on its verdict."
            )
            r.bold = True
            r.font.color.rgb = RGBColor.from_string("843C0C")
    else:
        p = doc.add_paragraph("No material issues found in the detector reasoning.")
        p.runs[0].italic = True

    _add_section_if_text(doc, "Recommended Action", 2,
                         judge_result.get("recommended_action", ""))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This audit complements (does not replace) human review.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(_FOOTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Report Generated: {today}    |    Analysis Tool: AI Content Detection System — LLM Judge")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(_FOOTER)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
