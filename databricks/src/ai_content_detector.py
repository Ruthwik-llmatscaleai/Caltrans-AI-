import json
import re
import datetime
import io
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from PyPDF2 import PdfReader
import docx as docx_lib
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.databricks_client import get_openai_client

CLAUDE_MODEL = "databricks-claude-sonnet-4-6"

_client = None


def _get_client():
    global _client
    if _client is None:
        _client = get_openai_client()
    return _client


def _extract_text_from_legacy_doc(data: bytes) -> str:
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".doc") as temp_file:
            temp_file.write(data)
            temp_path = Path(temp_file.name)

        antiword_path = shutil.which("antiword")
        if antiword_path:
            command = [antiword_path, str(temp_path)]
        elif sys.platform == "darwin" and shutil.which("textutil"):
            command = ["textutil", "-convert", "txt", "-stdout", str(temp_path)]
        else:
            raise ValueError("Legacy DOC extraction is not available in this runtime.")

        result = subprocess.run(command, check=True, capture_output=True, text=True)
        return result.stdout
    finally:
        if temp_path:
            temp_path.unlink(missing_ok=True)


def extract_text_from_file(file_obj, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        reader = PdfReader(file_obj)
        out = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                out += t + "\n"
        return out
    if name.endswith(".docx"):
        data = file_obj.read() if hasattr(file_obj, "read") else open(file_obj, "rb").read()
        d = docx_lib.Document(io.BytesIO(data))
        return "\n".join(p.text for p in d.paragraphs if p.text.strip())
    if name.endswith(".doc"):
        data = file_obj.read() if hasattr(file_obj, "read") else open(file_obj, "rb").read()
        return _extract_text_from_legacy_doc(data)
    raise ValueError("Unsupported file type. Upload a PDF, DOC, DOCX, or TXT.")


def split_sentences(text: str) -> list:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []
    parts = re.split(r"(?<=[.!?])\s+(?=[A-Z0-9\"'])", text)
    return [p.strip() for p in parts if p.strip()]


def _extract_json_object(raw: str) -> str:
    """Return the substring from the first '{' to its matching '}', ignoring
    braces that appear inside string literals. Falls back to the original
    string if no balanced object is found."""
    s = raw.strip()
    if s.startswith("```"):
        s = re.sub(r"^```(?:json)?\s*", "", s)
        s = re.sub(r"\s*```\s*$", "", s)
    start = s.find("{")
    if start < 0:
        return s
    depth = 0
    in_str = False
    esc = False
    for i in range(start, len(s)):
        c = s[i]
        if esc:
            esc = False
            continue
        if c == "\\":
            esc = True
            continue
        if c == '"':
            in_str = not in_str
            continue
        if in_str:
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                return s[start:i + 1]
    return s[start:]


def _repair_inner_quotes(s: str) -> str:
    """Escape unescaped double-quotes that appear inside JSON string values.
    Detects 'closing' quotes by peeking at the next non-whitespace char — a
    real string close is always followed by ',' ':' '}' or ']'."""
    out = []
    in_string = False
    esc = False
    i = 0
    n = len(s)
    while i < n:
        c = s[i]
        if esc:
            out.append(c)
            esc = False
            i += 1
            continue
        if c == "\\":
            out.append(c)
            esc = True
            i += 1
            continue
        if c == '"':
            if not in_string:
                in_string = True
                out.append(c)
            else:
                j = i + 1
                while j < n and s[j] in " \t\r\n":
                    j += 1
                if j >= n or s[j] in ",:}]":
                    in_string = False
                    out.append(c)
                else:
                    out.append('\\"')
            i += 1
            continue
        out.append(c)
        i += 1
    return "".join(out)


def _safe_json_parse(raw: str) -> dict:
    """Parse JSON from an LLM response, repairing the common failure modes:
    markdown fences, trailing commas, raw control chars in strings, and
    unescaped double-quotes inside string values. Raises json.JSONDecodeError
    only if every repair attempt fails."""
    payload = _extract_json_object(raw)
    try:
        return json.loads(payload)
    except json.JSONDecodeError:
        pass
    try:
        return json.loads(payload, strict=False)
    except json.JSONDecodeError:
        pass
    no_trailing = re.sub(r",(\s*[}\]])", r"\1", payload)
    try:
        return json.loads(no_trailing, strict=False)
    except json.JSONDecodeError:
        pass
    repaired = _repair_inner_quotes(no_trailing)
    return json.loads(repaired, strict=False)


SYSTEM_PROMPT = """You are an expert AI Content Detector analyzing a personal narrative for signs of AI generation. These narratives are written by small-business owners and individuals applying for state certifications, so many authors are non-native English speakers writing in a formal "application" register. Your job is to distinguish real AI fingerprints from formal-but-human writing — false positives have real consequences for applicants.

You will receive the full narrative AND a numbered list of every sentence. Produce:
1. A document-level verdict, using these definitions:
   - "Human Written": authored by a person. Light grammar/spelling polish (Grammarly, spell-check, an editor) is fine.
   - "AI-assisted": clear signs an LLM rewrote or expanded a human draft — e.g., one or two highly templated sections embedded inside otherwise specific personal content.
   - "AI-generated": pervasive AI fingerprints throughout, with little or no specific lived detail; reads like a model wrote it from a prompt.
2. Document-level percentages for Human / AI / Uncertain (must sum to 100).
3. Logical sections (Opening, Background, Barriers, Achievements, Closing, etc.) with per-section verdicts.
4. For EVERY sentence by id: verdict ("Human" | "AI" | "Uncertain"), score 0.0–1.0 (1.0 = strongly human), one short reason.
5. Concrete linguistic markers of human writing (with quoted evidence).
6. Concrete linguistic markers of AI writing (with quoted evidence and section/location).

============================
WHAT MODERN AI WRITING LOOKS LIKE (2024–2026)
============================
No single marker is decisive. Flag a sentence or section only when SEVERAL independent tells co-occur.

Lexical fingerprints — these words appear ~10–200x more often in AI than in human writing:
- Verbs: delve, leverage, foster, navigate, harness, underscore, embody, weave, illuminate, transcend, exemplify, streamline, elucidate, unlock, showcase.
- Nouns: tapestry, realm, landscape, journey, testament, interplay, nuance, symphony, paradigm, ecosystem, fabric, cornerstone, hallmark.
- Adjectives: pivotal, robust, comprehensive, meticulous, multifaceted, profound, seamless, holistic, vibrant, intricate, transformative, unparalleled, invaluable.
- Adverbs: meticulously, profoundly, seamlessly, intricately, vibrantly, notably, pivotally, tirelessly.
- Phrasal tells: "it is important to note", "plays a pivotal role", "in today's [X] landscape", "navigate the complexities of", "deeply rooted in", "stands as a testament to", "rich tapestry of", "in the realm of", "shed light on", "valuable insights".

Structural / stylistic tells:
- Em-dashes used for parenthetical or dramatic pauses, often more than once per page.
- Parallel triplets: "not only X, but also Y, and Z"; comma-separated lists of three modifier+noun phrases.
- Low burstiness — uniform paragraph length and uniform sentence length. Humans naturally mix 4-word sentences with 35-word ones.
- A "setup → body → recap" mini-arc inside a single paragraph.
- A closing paragraph that restates the prompt or recaps prior content with "In conclusion / Overall / In summary".
- Bullet points appearing mid-narrative in a personal essay.
- No contractions ("I have" / "do not" rather than "I've" / "don't") in otherwise informal-sounding voice.
- Grammar that is too clean: no fragments, no run-ons, no comma splices, no idiosyncratic punctuation, never opens a sentence with "And" or "But".
- Abstract emotional claims ("a profound sense of purpose", "this transformative experience") with no anchoring specifics.

Content tells (the strongest signal):
- Generic where specifics should exist: no proper nouns, no dollar amounts, no named places, no dated events, no named clients/teachers/family members.
- "Universal" struggles described in interchangeable terms — passages that could belong verbatim to any applicant's narrative.
- Logical scaffolding without lived texture: cause → effect → meaning, all neatly resolved.

============================
WHAT HUMAN WRITING LOOKS LIKE
============================
Any ONE of these is a meaningful human signal:
- Specific proper nouns: real names, towns, schools, employers, dates, dollar amounts, license numbers, project names.
- Idiosyncratic detail that doesn't serve the argument (a tangent, a remembered smell, a side-character).
- Quoted speech, dialect, untranslated foreign words, code-switching.
- Disfluencies, mid-sentence corrections, mixed register, casual asides.
- Personal voice: humor, anger, self-deprecation, embarrassment.
- Inconsistent capitalization/punctuation that reflects habit, not style.
- Sentence-length variance (mix of fragments and long sentences).
- Specific numbers tied to lived experience ("$4,200 in startup capital", "drove 47 miles each way").

============================
NEAR-DISPOSITIVE HUMAN SIGNALS (apply BEFORE any AI verdict)
============================
Two signal classes are nearly impossible to fake with current LLMs. If the document has BOTH at any density, the verdict is almost certainly Human Written, even when structural AI-looking features (parallel triplets, em-dashes, formal register, "Finally,…" closer) are also present:

1. Spelling errors and inconsistent punctuation. LLMs produce clean orthography. Misspellings ("Corrutption", "Kerela" for Kerala, "Gujrat" for Gujarat, "muncipalities"), mid-word hyphen artifacts ("stig-ma"), spurious ".:" or ",." sequences, inconsistent spaces around em-dashes, mixed curly/straight quotes — these are typed-by-a-human tells. Three or more such errors in a document of any length is a strong human signal. Grammarly-polish removes some of these; their presence proves the text was NOT run through such a tool, which itself argues against AI.
2. Dense, cited, attributed factual scaffolding. Real proper nouns clustered with sources, years, and statistics — e.g. "ADR reports 188 of 543 16th Lok Sabha members", "World Bank Doing Business Report (2016) ranks India 130/189", "Article 311", "Prevention of Corruption Act 1988". When a document sustains five or more such named/cited specifics across paragraphs, that is encyclopedic recall typical of a domain expert (UPSC aspirant, civil-service candidate, lawyer, academic), not LLM output to a generic prompt.

If both conditions are met, doc_verdict MUST be "Human Written" and confidence should reflect that the AI-looking surface features are policy/essayistic register, not generation artifacts. Do not soften to "AI-assisted" without evidence of a specific section that lacks both error-density and citation-density — and even then, name that section.

============================
FALSE-POSITIVE GUARDS — READ CAREFULLY
============================
Certification narratives are a templated genre. Several patterns LOOK like AI but are usually human:
- Non-native English (incl. South Asian / Indian English, Vietnamese English, Spanish-influenced English, etc.): simpler vocabulary, predictable phrasing, occasional article or preposition errors, calques from another language, regional idioms ("as under", "do the needful", "as is applicable", "the same shall be"). THIS IS NOT AI — peer-reviewed research shows detectors misclassify ~61% of non-native English essays as AI-generated. If you see consistent ESL patterns AND specific lived detail, lean Human.
- Grammarly / spell-check polish: a human draft cleaned up by a grammar tool is still Human Written. Clean grammar alone is not evidence of AI.
- Formal "application register" or essayistic register (UPSC essays, policy memos, op-eds): applicants and essayists sound formal because the genre demands it. Parallel triplets, em-dashes, and "Finally,…" closers are essay-genre conventions, not AI fingerprints. Formality + specifics = Human.
- Ghostwriter, family member, or accountant edited it: still Human authored from the applicant's lived experience.
- Classify as AI-assisted or AI-generated only when you see lexical AI fingerprints AND structural AI fingerprints AND a lack of lived specifics. All three are required, not preferred.

============================
RULES OF THE ROAD
============================
- Ground every claim in actual quoted text. Do NOT invent statistics like "12% passive voice" unless you can quote the exact sentences you counted.
- Confidence must reflect uncertainty. Reserve 0.85+ for cases where multiple independent fingerprints stack. Use 0.5–0.7 for plausible-but-not-conclusive.
- If the narrative is short (under 8 sentences), cap your confidence at 0.7 — there is too little signal.
- Per-sentence reasons: short and plain. "AI" or "Uncertain" verdicts must cite the specific tell (e.g., "lexical: 'tapestry' + 'pivotal'", "abstract; no specifics", "em-dash + parallel triplet").
- Personal, emotional, lived-experience phrasing is human, even when grammar is clean.
- Before issuing any AI or AI-assisted verdict, run the dispositive-human check above. If it triggers, you must explain in executive_summary why you overrode it (and you should rarely override it).
- When in doubt, prefer "Uncertain" over "AI". The cost of a false AI flag on a real applicant is high.

OUTPUT: valid JSON only, exactly this schema. No prose before or after.
{
  "doc_verdict": "Human Written" | "AI-assisted" | "AI-generated",
  "doc_confidence": 0.0,
  "summary_pct": {"human": 0, "ai": 0, "uncertain": 0},
  "sections": [
    {"name": "string", "verdict": "Human" | "AI-assisted" | "AI-generated" | "Mixed", "confidence": 0.0, "summary": "one sentence"}
  ],
  "sentences": [
    {"id": "s_1", "text": "exact sentence", "verdict": "Human" | "AI" | "Uncertain", "score": 0.0, "reason": "short"}
  ],
  "human_markers": [
    {"type": "string", "evidence": "short quote from the narrative"}
  ],
  "ai_markers": [
    {"type": "string", "evidence": "short quote from the narrative", "location": "section name or paragraph index"}
  ],
  "executive_summary": "2-4 sentence overall assessment in plain language",
  "methodology_notes": "brief sentence on how this was analyzed"
}"""


def detect_ai_content(narrative_text: str, max_tokens: int = 16000) -> dict:
    sentences = split_sentences(narrative_text)
    numbered = "\n".join(f"s_{i+1}: {s}" for i, s in enumerate(sentences))

    user_msg = (
        f"FULL NARRATIVE:\n{narrative_text}\n\n"
        f"NUMBERED SENTENCES (use these exact IDs in your output):\n{numbered}"
    )

    raw = ""
    try:
        resp = _get_client().chat.completions.create(
            model=CLAUDE_MODEL,
            max_tokens=max_tokens,
            temperature=0.0,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_msg},
            ],
        )
        raw = resp.choices[0].message.content.strip()
        return _safe_json_parse(raw)
    except json.JSONDecodeError as e:
        return {"error": f"Detector returned malformed JSON: {e}", "raw": raw[:1000]}
    except Exception as e:
        return {"error": f"Detector AI call failed: {e}"}


def _to_float(v, default=0.0):
    try:
        return float(v)
    except (TypeError, ValueError):
        return default


def generate_detector_report(result: dict, source_name: str = "Document") -> str:
    if "error" in result:
        return f"### AI Content Detection Report\n\n**Error:** {result['error']}\n"

    today = datetime.datetime.now().strftime("%B %d, %Y")
    pct = result.get("summary_pct", {}) or {}
    conf = _to_float(result.get("doc_confidence", 0.0))
    verdict = result.get("doc_verdict", "Unknown")

    md = "### AI Content Detection Report\n"
    md += "**Personal Narrative Analysis**\n\n"
    md += f"**Document:** {source_name}\n"
    md += f"**Date Analyzed:** {today}\n\n"

    md += "#### Executive Summary\n"
    md += result.get("executive_summary", "") + "\n\n"
    md += f"**Overall Assessment:** {verdict} | Confidence: {round(conf*100,1)}%\n\n"

    md += "#### Document-Level Summary\n\n"
    md += "| Human Written | Possibly AI | Uncertain | Confidence |\n"
    md += "|---|---|---|---|\n"
    md += f"| {pct.get('human', 0)}% | {pct.get('ai', 0)}% | {pct.get('uncertain', 0)}% | {round(conf*100,1)}% |\n\n"

    sections = result.get("sections", []) or []
    if sections:
        md += "#### Section-by-Section Analysis\n\n"
        md += "| Section | Classification | Confidence | Summary |\n"
        md += "|---|---|---|---|\n"
        for s in sections:
            sc = round(_to_float(s.get("confidence", 0.0)) * 100, 1)
            md += f"| {s.get('name','')} | {s.get('verdict','')} | {sc}% | {s.get('summary','')} |\n"
        md += "\n"

    sentences = result.get("sentences", []) or []
    flagged = [s for s in sentences if s.get("verdict") in ("AI", "Uncertain")]
    if flagged:
        md += "#### Flagged Sentences\n\n"
        md += "| ID | Verdict | Score | Sentence | Reason |\n"
        md += "|---|---|---|---|---|\n"
        for s in flagged:
            txt = s.get("text", "").replace("|", "\\|")
            reason = s.get("reason", "").replace("|", "\\|")
            md += f"| {s.get('id','')} | {s.get('verdict','')} | {round(_to_float(s.get('score',0)),2)} | {txt} | {reason} |\n"
        md += "\n"
    else:
        md += "#### Flagged Sentences\n\nNone — all sentences read as human-written.\n\n"

    hm = result.get("human_markers", []) or []
    if hm:
        md += "#### Human Writing Markers Present\n\n"
        md += "| Marker Type | Evidence |\n|---|---|\n"
        for m in hm:
            ev = (m.get("evidence", "") or "").replace("|", "\\|")
            md += f"| {m.get('type','')} | {ev} |\n"
        md += "\n"

    am = result.get("ai_markers", []) or []
    if am:
        md += "#### AI Writing Markers\n\n"
        md += "| Marker Type | Evidence | Location |\n|---|---|---|\n"
        for m in am:
            ev = (m.get("evidence", "") or "").replace("|", "\\|")
            loc = (m.get("location", "") or "").replace("|", "\\|")
            md += f"| {m.get('type','')} | {ev} | {loc} |\n"
        md += "\n"

    md += "#### Methodology Notes\n"
    md += result.get("methodology_notes", "") + "\n\n"
    md += "*This is an advisory tool. AI detection is probabilistic, not definitive.*\n"

    return md.replace("$", "\\$")


# ============================================================
# Word document generator. The LLM produces only the structured
# `result` JSON; all design (fonts, colors, borders, layout) is
# fixed below so the report looks the same every run.
# ============================================================

_PRIMARY = "1F4E79"
_HEADER_FILL = "1F4E79"
_HEADER_TEXT = "FFFFFF"
_ZEBRA_FILL = "F7F9FC"
_BORDER_COLOR = "BFBFBF"
_MUTED = "595959"
_FOOTER = "808080"

_VERDICT_COLORS = {
    "human": ("E2EFDA", "375623"),
    "ai-assisted": ("FFF2CC", "806000"),
    "ai-generated": ("FCE4D6", "843C0C"),
    "ai": ("FCE4D6", "843C0C"),
    "uncertain": ("FFF2CC", "806000"),
    "mixed": ("FFF2CC", "806000"),
}

_DOCX_UNSAFE_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\uD800-\uDFFF\uFFFE\uFFFF]")


def _verdict_palette(verdict: str):
    v = (verdict or "").strip().lower()
    if "human" in v:
        return _VERDICT_COLORS["human"]
    if "ai-generated" in v or v == "ai":
        return _VERDICT_COLORS["ai-generated"]
    if "ai-assisted" in v:
        return _VERDICT_COLORS["ai-assisted"]
    if "mixed" in v:
        return _VERDICT_COLORS["mixed"]
    if "uncertain" in v:
        return _VERDICT_COLORS["uncertain"]
    return ("F2F2F2", _MUTED)


def _shade_cell(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _clean_docx_text(text) -> str:
    return _DOCX_UNSAFE_RE.sub("", str(text or ""))


def _set_cell(cell, text, *, bold=False, color_hex=None, size=10, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(_clean_docx_text(text))
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    if color_hex:
        r.font.color.rgb = RGBColor.from_string(color_hex)


def _set_table_borders(tbl, color=_BORDER_COLOR, size="6"):
    tbl_pr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tbl_pr.append(borders)


def _cm_to_dxa(cm: float) -> int:
    return int(round(cm * 567))


def _set_col_widths(tbl, widths_cm):
    # Force fixed layout so Word honors the widths instead of autofit.
    tbl_pr = tbl._tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblLayout")):
        tbl_pr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(_cm_to_dxa(w) for w in widths_cm)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)

    for existing in tbl._tbl.findall(qn("w:tblGrid")):
        tbl._tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(_cm_to_dxa(w)))
        grid.append(gc)
    tbl._tbl.insert(list(tbl._tbl).index(tbl_pr) + 1, grid)

    for row in tbl.rows:
        # Prevent a single row from being split across a page break.
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _add_header_row(tbl, headers, *, aligns=None, size=10):
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _shade_cell(c, _HEADER_FILL)
        a = aligns[i] if aligns and i < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell(c, h, bold=True, color_hex=_HEADER_TEXT, size=size, align=a)


_NBSP = " "


def _no_widow(text: str) -> str:
    """Replace the last space in heading text with a non-breaking space so the
    last word can't end up alone on its own line (typography widow fix)."""
    s = (text or "").rstrip()
    idx = s.rfind(" ")
    if idx < 0:
        return s
    return s[:idx] + " " + s[idx + 1:]


def _style_heading(paragraph, color_hex=_PRIMARY):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(color_hex)
        run.font.name = "Calibri"


def _add_styled_heading(doc, text, level, color_hex=_PRIMARY):
    h = doc.add_heading(_no_widow(_clean_docx_text(text)), level=level)
    _style_heading(h, color_hex)
    return h


def _fmt_pct(value_0_to_1, decimals=1):
    pct = round(float(value_0_to_1) * 100, decimals)
    if pct == int(pct):
        return f"{int(pct)}%"
    return f"{pct}%"


def _add_section_if_text(doc, heading, level, text, color_hex=_PRIMARY):
    body = _clean_docx_text(text).strip()
    if not body:
        return False
    _add_styled_heading(doc, heading, level=level, color_hex=color_hex)
    doc.add_paragraph(body)
    return True


def generate_detector_docx(result: dict, source_name: str = "Document") -> bytes:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    today = datetime.datetime.now().strftime("%B %d, %Y")

    if "error" in result:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("AI Content Detection Report")
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = RGBColor.from_string(_PRIMARY)
        doc.add_paragraph(f"Error: {result['error']}")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    pct = result.get("summary_pct", {}) or {}
    conf = _to_float(result.get("doc_confidence", 0.0))
    verdict = result.get("doc_verdict", "Unknown")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AI Content Detection Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string(_PRIMARY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Personal Narrative Analysis")
    r.italic = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(_MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Document: {source_name}    |    Date Analyzed: {today}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(_MUTED)

    _add_section_if_text(doc, "Executive Summary", 1,
                         result.get("executive_summary", ""))

    badge_p = doc.add_paragraph()
    fill, text_color = _verdict_palette(verdict)
    r1 = badge_p.add_run("Overall Assessment: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = badge_p.add_run(verdict)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor.from_string(text_color)
    r3 = badge_p.add_run(f"   |   Confidence: {_fmt_pct(conf)}")
    r3.font.size = Pt(11)

    _add_styled_heading(doc, "Document-Level Summary", 2)
    tbl = doc.add_table(rows=2, cols=4)
    centered = [WD_ALIGN_PARAGRAPH.CENTER] * 4
    _add_header_row(tbl, ["Human Written", "Possibly AI", "Uncertain", "Confidence"],
                    aligns=centered, size=11)
    vals = [
        f"{pct.get('human', 0)}%",
        f"{pct.get('ai', 0)}%",
        f"{pct.get('uncertain', 0)}%",
        _fmt_pct(conf),
    ]
    for i, v in enumerate(vals):
        _set_cell(tbl.rows[1].cells[i], v, size=11, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [4.0, 4.0, 4.0, 4.0])

    sections = result.get("sections", []) or []
    if sections:
        _add_styled_heading(doc, "Section-by-Section Analysis", 2)
        tbl = doc.add_table(rows=1 + len(sections), cols=4)
        _add_header_row(tbl, ["Section", "Classification", "Confidence", "Summary"],
                        aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
        for ri, s in enumerate(sections, start=1):
            v_str = s.get("verdict", "") or ""
            fill, text_color = _verdict_palette(v_str)
            row = tbl.rows[ri]
            _set_cell(row.cells[0], s.get("name", "") or "", size=10, bold=True)
            _set_cell(row.cells[1], v_str, size=10, bold=True, color_hex=text_color)
            _shade_cell(row.cells[1], fill)
            _set_cell(row.cells[2], _fmt_pct(s.get("confidence", 0.0)),
                      size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(row.cells[3], s.get("summary", "") or "", size=10)
            if ri % 2 == 0:
                for ci in (0, 2, 3):
                    _shade_cell(row.cells[ci], _ZEBRA_FILL)
        _set_table_borders(tbl)
        _set_col_widths(tbl, [4.5, 3.5, 2.5, 6.0])

    _add_styled_heading(doc, "Flagged Sentences", 2)
    sentences = result.get("sentences", []) or []
    flagged = [s for s in sentences if s.get("verdict") in ("AI", "Uncertain")]
    if flagged:
        tbl = doc.add_table(rows=1 + len(flagged), cols=5)
        _add_header_row(tbl, ["ID", "Verdict", "Score", "Sentence", "Reason"],
                        aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                                WD_ALIGN_PARAGRAPH.LEFT])
        for ri, s in enumerate(flagged, start=1):
            v_str = s.get("verdict", "") or ""
            fill, text_color = _verdict_palette(v_str)
            row = tbl.rows[ri]
            _set_cell(row.cells[0], s.get("id", "") or "", size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(row.cells[1], v_str, size=9, bold=True, color_hex=text_color,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade_cell(row.cells[1], fill)
            _set_cell(row.cells[2], f"{round(_to_float(s.get('score', 0)), 2)}",
                      size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell(row.cells[3], s.get("text", "") or "", size=9)
            _set_cell(row.cells[4], s.get("reason", "") or "", size=9)
            if ri % 2 == 0:
                for ci in (0, 2, 3, 4):
                    _shade_cell(row.cells[ci], _ZEBRA_FILL)
        _set_table_borders(tbl)
        _set_col_widths(tbl, [1.5, 2.0, 1.5, 7.0, 4.5])
    else:
        p = doc.add_paragraph("None — all sentences read as human-written.")
        p.runs[0].italic = True

    hm = result.get("human_markers", []) or []
    am = result.get("ai_markers", []) or []
    if hm or am:
        _add_styled_heading(doc, "Linguistic Analysis", 2)
    if hm:
        _add_styled_heading(doc, "Human Writing Markers Present", 3)
        tbl = doc.add_table(rows=1 + len(hm), cols=2)
        _add_header_row(tbl, ["Marker Type", "Evidence"])
        for ri, m in enumerate(hm, start=1):
            row = tbl.rows[ri]
            _set_cell(row.cells[0], m.get("type", "") or "", size=10, bold=True)
            _set_cell(row.cells[1], m.get("evidence", "") or "", size=10)
            if ri % 2 == 0:
                for ci in (0, 1):
                    _shade_cell(row.cells[ci], _ZEBRA_FILL)
        _set_table_borders(tbl)
        _set_col_widths(tbl, [4.5, 12.0])
    if am:
        _add_styled_heading(doc, "AI Writing Markers", 3)
        tbl = doc.add_table(rows=1 + len(am), cols=3)
        _add_header_row(tbl, ["Marker Type", "Evidence", "Location"])
        for ri, m in enumerate(am, start=1):
            row = tbl.rows[ri]
            _set_cell(row.cells[0], m.get("type", "") or "", size=10, bold=True)
            _set_cell(row.cells[1], m.get("evidence", "") or "", size=10)
            _set_cell(row.cells[2], m.get("location", "") or "", size=10)
            if ri % 2 == 0:
                for ci in (0, 1, 2):
                    _shade_cell(row.cells[ci], _ZEBRA_FILL)
        _set_table_borders(tbl)
        _set_col_widths(tbl, [3.5, 9.5, 3.5])

    _add_section_if_text(doc, "Methodology Notes", 2,
                         result.get("methodology_notes", ""))

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("This is an advisory tool. AI detection is probabilistic, not definitive.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(_FOOTER)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Report Generated: {today}    |    Analysis Tool: AI Content Detection System")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(_FOOTER)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
